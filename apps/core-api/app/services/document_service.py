import os
import sys
import uuid
from typing import List, Optional
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from openai import OpenAI
from sqlalchemy.orm import Session

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))
from phi_utils.retry import retry_async
from phi_utils.logging import setup_logging

logger = setup_logging("core-api.document_service")

from app.config import settings
from app.models import Document, DocumentChunk

# Initialize OpenAI client - lazy initialization to ensure settings are loaded
def get_openai_client():
    """Get OpenAI client with current API key from settings"""
    api_key = settings.openai_api_key or os.getenv("OPENAI_API_KEY", "")
    if not api_key or api_key == "your-openai-api-key-here":
        raise ValueError("OPENAI_API_KEY is not set or is still the placeholder value")
    return OpenAI(api_key=api_key)


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract text from various file formats"""
    ext = file_extension.lower()
    
    if ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    
    elif ext in [".docx", ".doc"]:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
    """Split text into chunks with overlap"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence boundary
        if end < len(text):
            last_period = chunk.rfind(".")
            last_newline = chunk.rfind("\n")
            break_point = max(last_period, last_newline)
            if break_point > chunk_size * 0.5:  # Only break if we're past halfway
                chunk = chunk[:break_point + 1]
                end = start + break_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap  # Overlap for context
    
    return chunks


def generate_embedding(text: str) -> List[float]:
    """Generate embedding using OpenAI with retry (sync version)"""
    from phi_utils.retry import retry_sync
    
    def call_embedding():
        client = get_openai_client()
        return client.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
    
    try:
        response = retry_sync(
            call_embedding,
            max_retries=3,
            delay=0.5,
            backoff=2.0,
            exceptions=(Exception,),
            logger=logger
        )
        return response.data[0].embedding
    except Exception as e:
        logger.error(f"Failed to generate embedding after retries: {str(e)}")
        raise


def process_document(
    db: Session,
    file_path: str,
    file_name: str,
    org_id: uuid.UUID,
    agent_id: Optional[uuid.UUID],
    source_type: str
) -> Document:
    """Process uploaded document: extract, chunk, embed, and store"""
    
    # Extract file extension
    file_ext = Path(file_name).suffix
    
    # Extract text
    text = extract_text_from_file(file_path, file_ext)
    
    if not text.strip():
        raise ValueError("No text could be extracted from the document")
    
    # Create document record
    document = Document(
        org_id=org_id,
        agent_id=agent_id,
        name=file_name,
        source_type=source_type,
        storage_path=file_path
    )
    db.add(document)
    db.flush()
    
    # Chunk text
    chunks = chunk_text(text)
    
    # Generate embeddings and store chunks
    for i, chunk_text_content in enumerate(chunks):
        embedding = generate_embedding(chunk_text_content)
        
        # Convert embedding to string format for pgvector
        embedding_str = "[" + ",".join(map(str, embedding)) + "]"
        
        # Use raw SQL to insert with vector type
        from sqlalchemy import text
        import json
        metadata_json = json.dumps({
            "chunk_index": i,
            "total_chunks": len(chunks),
            "source_type": source_type
        })
        db.execute(
            text("""
                INSERT INTO document_chunks 
                (id, document_id, org_id, agent_id, chunk_text, embedding, metadata, created_at)
                VALUES 
                (uuid_generate_v4(), :doc_id, :org_id, :agent_id, :chunk_text, :embedding::vector, :metadata::jsonb, now())
            """),
            {
                "doc_id": str(document.id),
                "org_id": str(org_id),
                "agent_id": str(agent_id) if agent_id else None,
                "chunk_text": chunk_text_content,
                "embedding": embedding_str,
                "metadata": metadata_json
            }
        )
    
    db.commit()
    db.refresh(document)
    
    return document

